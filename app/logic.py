import os
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

env_path = Path(__file__).resolve().parent.parent / ".env"
load_dotenv(dotenv_path=env_path)

api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

def extract_text_from_pdf(file_path: str):
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def extract_text_from_docx(file_path: str):
    """ კითხულობს ტექსტს პარაგრაფებიდან და ცხრილებიდან """
    try:
        doc = Document(file_path)
        full_text = []
        
        # 1. ტექსტი პარაგრაფებიდან
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # 2. ტექსტი ცხრილებიდან (CV-ებისთვის აუცილებელია!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def extract_text_from_file(file_path: str):
    # ვაქცევთ პატარა ასოებად (.DOCX -> .docx), რომ არ აირიოს
    path_lower = file_path.lower()
    
    if path_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif path_lower.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        return None

def analyze_cv(cv_text: str, job_description: str):
    # თუ ტექსტი ძალიან მოკლეა, ესე იგი ვერ წაიკითხა
    if len(cv_text.strip()) < 10:
        return '{"error": "ფაილი წაიკითხა, მაგრამ ტექსტი ძალიან ცოტაა. სცადეთ სხვა ფაილი."}'

    prompt = f"""
    You are an expert HR AI Assistant. 
    Compare the following Candidate CV with the Job Description.
    
    JOB DESCRIPTION:
    {job_description}
    
    CANDIDATE CV:
    {cv_text}
    
    Provide a response in strict JSON format with the following keys:
    - match_score: (integer between 0-100)
    - missing_skills: (list of strings, technical skills from JD that are missing in CV)
    - summary: (string, brief explanation of the decision)
    - recommendation: (string, advice for the candidate)
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful HR assistant that outputs JSON."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error connecting to AI: {e}"